"""Generate a polished Word document for the Travelers Group sales lead brief."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Travelers_Lead_Brief.docx"

# EXL brand palette
NAVY = RGBColor(0x0B, 0x2C, 0x4E)
ACCENT = RGBColor(0xC8, 0x10, 0x2E)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY = RGBColor(0xEE, 0xEE, 0xEE)


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_heading(p, text, size_pt=18, color=NAVY, bold=True, space_before=12, space_after=6):
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)


def add_para(doc, text, size_pt=10.5, bold=False, italic=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, size_pt=10.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_kv_table(doc, rows, label_width=Inches(2.0), value_width=Inches(4.5)):
    """Two-column key-value table."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(rows):
        kc, vc = tbl.rows[i].cells
        kc.width = label_width
        vc.width = value_width
        # label
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.font.name = "Arial"
        kr.font.size = Pt(10)
        kr.font.bold = True
        kr.font.color.rgb = NAVY
        # value
        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.name = "Arial"
        vr.font.size = Pt(10)
        shade_cell(kc, "F5F7FA")
        set_cell_borders(kc)
        set_cell_borders(vc)
    return tbl


def add_data_table(doc, headers, rows, col_widths=None):
    """Generic data table with header row + body rows."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "0B2C4E")
        set_cell_borders(hdr[i], color="0B2C4E")
        if col_widths:
            hdr[i].width = col_widths[i]
    # body
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            set_cell_borders(cells[ci])
            if col_widths:
                cells[ci].width = col_widths[ci]
            if ri % 2 == 1:
                shade_cell(cells[ci], "F5F7FA")
    return tbl


# ---------------- Build document ----------------
doc = Document()

# Page setup: US Letter, 1" margins
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)

# Default style
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10.5)

# ===== TITLE BLOCK =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("EXL Sales Lead Brief")
tr.font.name = "Arial"
tr.font.size = Pt(11)
tr.font.color.rgb = ACCENT
tr.font.bold = True
title.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
sr = sub.add_run("Travelers Group")
sr.font.name = "Arial"
sr.font.size = Pt(28)
sr.font.bold = True
sr.font.color.rgb = NAVY
sub.paragraph_format.space_after = Pt(2)

tag = doc.add_paragraph()
tag_r = tag.add_run("Lead #1 of 236  •  Score 89.3 / 100  •  Confidence: Very High  •  17 triggers fired (9 unique types)")
tag_r.font.name = "Arial"
tag_r.font.size = Pt(10)
tag_r.font.italic = True
tag_r.font.color.rgb = GREY
tag.paragraph_format.space_after = Pt(14)

# divider rule
rule = doc.add_paragraph()
p_pr = rule._p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:color"), "C8102E")
p_bdr.append(bottom)
p_pr.append(p_bdr)
rule.paragraph_format.space_after = Pt(8)

# ===== EXECUTIVE SUMMARY =====
style_heading(doc.add_paragraph(), "Executive Summary", 16)
add_para(
    doc,
    "Travelers Group is the highest-scoring account in the TAM corpus by a wide margin "
    "(next account scores 64.5). The lead is driven by a rare alignment of three displacement "
    "signals at a single account: a mass renewal of ~1,000 Cognizant FTEs in <6 months, three "
    "incumbents simultaneously in champion-challenger friction, and a combined competitor "
    "footprint of ~1,210 FTEs across 9 engagements. A sequenced displacement could move "
    "multiple scopes within 12-18 months.",
    space_after=8,
)

# ===== ACCOUNT SNAPSHOT =====
style_heading(doc.add_paragraph(), "Account Snapshot", 14)
add_kv_table(doc, [
    ("Account", "Travelers Group"),
    ("Type / Geography", "Carrier / US"),
    ("Size", "$27.2B NWP 2019 (+5.5% YoY)  |  $29B GWP"),
    ("Rank by NWP", "#6"),
    ("Combined ratio", "96.5"),
    ("EXL relationship", "Named  (EXL Client column: blank — verify)"),
    ("Outsources", "Yes"),
    ("EXL Growth Leader", "Sean"),
    ("EXL Client Exec", "Sean"),
    ("Displacement owner", "Saju"),
])
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== SCORE BREAKDOWN =====
style_heading(doc.add_paragraph(), "Score Breakdown", 14)
add_data_table(
    doc,
    headers=["Component", "Score (0–1)", "Weight", "Contribution"],
    rows=[
        ["Account fit (NWP rank, Strategic/Named, growth)", "0.85", "30%", "0.255"],
        ["Signal strength (trigger sum + diversity)", "1.00", "40%", "0.400"],
        ["Solution match (triggers → EXL solutions)", "0.79", "20%", "0.158"],
        ["Relationship warmth (CE / Growth Leader)", "0.80", "10%", "0.080"],
        ["TOTAL", "", "", "0.893  (89.3 / 100)"],
    ],
    col_widths=[Inches(3.8), Inches(1.0), Inches(0.9), Inches(1.2)],
)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== TRIGGERS =====
style_heading(doc.add_paragraph(), "Triggers Firing (17 total)", 14)
trigger_rows = [
    ["1", "renewal_le_6mo", "1.00", "Cognizant contract renews in <6 months (~1000 FTEs across 5 scopes: Claims, UW, F&A)"],
    ["2", "mass_renewal", "1.00", "Single-vendor displacement window: Cognizant holds 5 concurrent scopes (~1000 FTEs) all renewing in <6 months"],
    ["3", "multi_competitor_friction", "1.00", "3 concurrent incumbents in friction: Cognizant, Genpact (terminating), WNS (given notice)"],
    ["4", "active_friction (Cognizant)", "0.90", "Champion-challenger across Personal Insurance, Commercial Insurance, claims settlement, Bill/Utilization Review"],
    ["5", "active_friction (Genpact)", "0.90", "Champion-challenger, terminating — Policy Maintenance, Claims, Invoicing"],
    ["6", "active_friction (WNS)", "0.90", "Given notice, exploring new champion"],
    ["7", "function_headroom (Cognizant)", "0.73", "Parked at Policy Servicing (~1000 FTEs); EXL retention 35% at this function → ~2,857 FTE TAM"],
    ["8", "function_headroom (Genpact)", "0.73", "Parked at Policy Servicing (~210 FTEs); ~600 FTE TAM at this function"],
    ["9", "high_competitor_ftes (Cognizant)", "0.70", "~1000 FTEs at a Named account — large displacement target"],
    ["10", "high_competitor_ftes (Genpact)", "0.70", "~210 FTEs at a Named account — large displacement target"],
    ["11", "unaddressed_voc", "0.60", "VOC: Optimizing Operating Leverage → 2 components mapped to EXL Int_Col, status Not Initiated"],
    ["12", "unaddressed_voc", "0.60", "VOC: Strong Top Line Performance → 3 components for EXL NB_PaaS, Not Initiated"],
    ["13", "strategic_low_wallet", "0.60", "Named account but EXL share-of-wallet = blank / EXL Client = blank"],
    ["14", "unaddressed_voc", "0.60", "VOC: Optimizing Operating Leverage → 4 components for EXL FNOL, Not Initiated"],
    ["15", "competitor_concentration", "0.60", "Aggregate competitor footprint ~1,210 FTEs (3 competitors, 9 engagements)"],
    ["16", "unaddressed_voc", "0.60", "VOC: Optimizing Operating Leverage → 1 component for EXL E_E2E, Not Initiated"],
    ["17", "unaddressed_voc", "0.60", "VOC: Strong Top Line Performance → 1 component for EXL FNOL, Not Initiated"],
]
add_data_table(
    doc,
    headers=["#", "Trigger", "Strength", "Evidence"],
    rows=trigger_rows,
    col_widths=[Inches(0.3), Inches(1.6), Inches(0.7), Inches(4.3)],
)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== RECOMMENDED EXL PLAY =====
style_heading(doc.add_paragraph(), "Recommended EXL Play", 14)
add_para(doc, "Anchor stack (most-cited across this account's triggers):", bold=True, space_after=6)
add_data_table(
    doc,
    headers=["EXL Product", "What it does", "Typical outcome"],
    rows=[
        ["EXL XTRAKTO.AI", "Intelligent document processing — automated data extraction from forms, ACORD, mail, email", "25-30% cost-of-operations reduction; 30% cycle-time improvement"],
        ["EXL EXELIA.AI", "Conversational AI — voicebot/chatbot for self-service", "AHT reduction; 5-10 pt NPS lift; higher First Call Resolution"],
        ["EXL NerveHub", "Workflow orchestration — connects intake, decisioning, ops across the value chain", "Cross-process efficiency; integrated end-to-end visibility"],
        ["EXL Paymentor", "Payment automation — AP, payment processing, leakage detection", "100% payment-leakage reduction; 40-60% F&A cost reduction"],
    ],
    col_widths=[Inches(1.6), Inches(3.4), Inches(1.9)],
)
doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_para(doc, "Primary trigger driving the pitch:", bold=True, space_after=2)
add_para(doc, "Incumbent renewal within 6 months. Solution mapping basis: functions Claims, UW, F&A in evidence.", space_after=6)
add_para(doc, "Secondary trigger:", bold=True, space_after=2)
add_para(doc, "Active competitor friction (free-text signal) — function Claims in evidence.", space_after=6)
add_para(doc, "All matched solutions for this account:", bold=True, space_after=2)
add_para(doc, "EXL XTRAKTO.AI (5)  •  EXL Paymentor (4)  •  EXL EXELIA.AI (4)  •  EXL Subrosource (3)  •  EXL MedConnection (3)", space_after=8)

# ===== PITCH HOOK =====
style_heading(doc.add_paragraph(), "Pitch Hook", 14)
hook = doc.add_paragraph()
hook_pr = hook._p.get_or_add_pPr()
left_border = OxmlElement("w:pBdr")
left = OxmlElement("w:left")
left.set(qn("w:val"), "single")
left.set(qn("w:sz"), "24")
left.set(qn("w:space"), "8")
left.set(qn("w:color"), "C8102E")
left_border.append(left)
hook_pr.append(left_border)
hook_run = hook.add_run(
    "Sequenced displacement vehicle for Cognizant's 1,000+ FTE footprint as multiple concurrent "
    "scopes approach renewal. Lead with EXL XTRAKTO.AI + EXL EXELIA.AI — typical outcome: "
    "25-30% reduction in cost of operations, 30% improvement in cycle time. Anchor the narrative "
    "on Travelers' stated operating-leverage commitments, with EXL's solution mapping to their "
    "VOC already complete."
)
hook_run.font.name = "Arial"
hook_run.font.size = Pt(11)
hook_run.font.italic = True
hook_run.font.color.rgb = NAVY
hook.paragraph_format.left_indent = Inches(0.15)
hook.paragraph_format.space_after = Pt(10)

# ===== STAKEHOLDERS =====
style_heading(doc.add_paragraph(), "Stakeholder Map (Client Side)", 14)
add_data_table(
    doc,
    headers=["Role", "Name"],
    rows=[
        ["CEO", "Alan Schnitzer"],
        ["CFO", "Dan Frey"],
        ["CRO", "Bruce Jones"],
        ["CCO", "Nicolas Seminara"],
        ["COO", "Mojgan Lefebvre"],
        ["CIO / CTO", "Mojgan Lefebvre"],
        ["CUO", "Rick Schug"],
    ],
    col_widths=[Inches(1.8), Inches(5.0)],
)
add_para(doc, "Note: COO and CIO/CTO listed under the same name — possible dual role or data-entry artifact. Verify before outreach.", italic=True, color=GREY, size_pt=9.5, space_after=10)

# ===== RISKS =====
style_heading(doc.add_paragraph(), "Risks & Disqualifiers", 14)
risks = [
    "EXL Client column is blank — confirm whether EXL has any current scope at this account before pitching.",
    "Data vintage 2020-2021 (~4-5 years old). CxO names and renewal windows may be stale — verify externally.",
    "Large incumbent footprint = high switching cost. Plan for a sequenced / phased displacement, not big-bang.",
    "Three parallel champion-challenger threads — coordinate internally to avoid fragmented or conflicting pitches.",
    "Unaddressed VOC items present but no documented current EXL scope — we are pitching from outside the wallet.",
    "COO and CIO/CTO show the same name — could be a real dual role or a source-data artifact. Verify.",
]
for r in risks:
    add_bullet(doc, r)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== NEXT 3 STEPS =====
style_heading(doc.add_paragraph(), "Next 3 Steps", 14)
add_data_table(
    doc,
    headers=["#", "Owner", "Action"],
    rows=[
        ["1", "Saju", "Engage primary client contact on champion-challenger discussions; bring case studies for the largest incumbent's scope."],
        ["2", "Saju", "Sequence the displacement timeline — smallest scope as fast-pilot, anchor scope as headline."],
        ["3", "Sean", "Schedule meeting with client CxO to walk through the EXL solution components already mapped to their stated priorities."],
    ],
    col_widths=[Inches(0.4), Inches(0.9), Inches(5.5)],
)
doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== EVIDENCE PACK =====
style_heading(doc.add_paragraph(), "Evidence Pack — Source Citations", 14)
add_para(doc, "Every trigger above is sourced from rows in the consolidated TAM corpus:", space_after=4)
evidence = [
    ("F5 / Competitor Analysis / R8–R12", "Cognizant footprint, scopes, renewal window, FTEs (5 rows)"),
    ("F5 / Competitor Analysis / R39–R41", "Genpact footprint, terminating signal, Policy Maintenance/Claims/Invoicing (3 rows)"),
    ("F5 / Competitor Analysis / R65", "WNS — given notice, exploring new champion"),
    ("F1 / Buyer Priorities / R9, R11–R20", "VOC priorities mapped to EXL solutions, status Not Initiated"),
    ("F1 / Top Insurers and Brokers", "Client profile, NWP, EXL relationship type, share-of-wallet flag"),
    ("Derived: F4 + F5", "Aggregate competitor footprint summary (~1,210 FTEs, 3 competitors, 9 engagements)"),
]
add_data_table(doc, headers=["Source", "What it provides"], rows=evidence, col_widths=[Inches(2.8), Inches(4.0)])
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ===== FOOTER =====
footer = doc.add_paragraph()
fr = footer.add_run("Generated 2026-05-27 from data vintages 2020-08-03 and 2021-02-21. Source-cited to TAM corpus (5 internal EXL spreadsheets). Re-run pipeline/synthesize.py to refresh.")
fr.font.name = "Arial"
fr.font.size = Pt(8.5)
fr.font.italic = True
fr.font.color.rgb = GREY
footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
footer.paragraph_format.space_before = Pt(12)

doc.save(OUT)
print(f"Wrote {OUT}")
